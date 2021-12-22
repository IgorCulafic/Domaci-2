""" from docx import Document

document = Document('demo.docx')
x = 0
for i in document.paragraphs:
    if x < 20:
        print(i.text)
    else:
        break;
    x = x + 1
for para in document.paragraphs:
    print(para.text) """
    
    
# string_x = "abc"
# data = ""
# for i in range (0, len(string_x)):
#     data += str(string_x[i])
# print(data)
    
from tkinter import *
import docx2txt
my_text = docx2txt.process("Carobno samarce - lat.docx")

#document = Document('Carobno samarce - lat.docx')
#Create window object
app = Tk()
# rijeci = ''
# x = 0
# c = 0
# page =6

# for i in document.paragraphs:
#     if x < page:
#         show = i.text
#         part_text = StringVar()
#         part_label = Label(app, text = show, font = ('bold',14), width = 50,wraplength=500, justify="center")
#         part_label.grid(row = c, column = 5,columnspan= 6)
        
#     else:
#         page += page
#         break;
#     x = x + 1
#     c = c + 1
#part

part_text = StringVar()
part_label = Label(app, text = my_text, font = ('bold',14), width = 50,wraplength=500, justify="center")
part_label.grid(row = 1, column = 5,columnspan= 6)

#Create Scrollbar
scrollbar = Scrollbar(app)
scrollbar.grid(row = 3, column = 3)

#Set scroll to listbox
part_label.configure(yscrollcommand = scrollbar.set)
scrollbar.configure(command = part_label.yview)


# part_label = Label(app, text = rijeci, font = ('bold',14), pady = 20)
# part_label.grid(row = 0, column = 0,rowspan= 2, sticky = W+E)

app.title('Domaci 2')
app.geometry('550x850')

#start program
app.mainloop()